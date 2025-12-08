# core/exporters.py
from __future__ import annotations

import logging
from dataclasses import dataclass
from io import BytesIO
from typing import Optional

logger = logging.getLogger("essay_writer.exporters")


class ExportError(RuntimeError):
    pass


@dataclass
class ExportBundle:
    md: bytes
    txt: bytes
    docx: bytes
    pdf: bytes


def _normalize_text(text: str) -> str:
    return (text or "").replace("\r\n", "\n").strip()


def export_markdown(text: str) -> bytes:
    return _normalize_text(text).encode("utf-8")


def export_txt(text: str) -> bytes:
    return _normalize_text(text).encode("utf-8")


def export_docx(title: str, essay_text: str) -> bytes:
    """
    Simple DOCX export using python-docx: title + paragraphs.
    python-docx supports Document() and add_paragraph(). :contentReference[oaicite:3]{index=3}
    """
    try:
        from docx import Document
    except Exception as e:
        raise ExportError(f"python-docx not installed/available: {e}") from e

    try:
        doc = Document()
        t = (title or "Essay").strip()
        doc.add_heading(t, level=1)

        text = _normalize_text(essay_text)
        for para in text.split("\n\n"):
            p = para.strip()
            if p:
                doc.add_paragraph(p)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        logger.exception("DOCX export failed")
        raise ExportError(f"DOCX export failed: {e}") from e


def export_pdf(title: str, essay_text: str) -> bytes:
    """
    PDF export using ReportLab Platypus (handles wrapping nicely).
    ReportLab supports building PDFs from Paragraph/Spacer via SimpleDocTemplate. :contentReference[oaicite:4]{index=4}
    """
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except Exception as e:
        raise ExportError(f"reportlab not installed/available: {e}") from e

    try:
        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=LETTER)
        styles = getSampleStyleSheet()

        story = []
        t = (title or "Essay").strip()
        story.append(Paragraph(t, styles["Title"]))
        story.append(Spacer(1, 12))

        text = _normalize_text(essay_text)
        for para in text.split("\n\n"):
            p = para.strip()
            if p:
                story.append(Paragraph(p.replace("\n", "<br/>"), styles["BodyText"]))
                story.append(Spacer(1, 10))

        doc.build(story)
        return buf.getvalue()
    except Exception as e:
        logger.exception("PDF export failed")
        raise ExportError(f"PDF export failed: {e}") from e


def build_export_bundle(*, title: str, essay_text: str) -> ExportBundle:
    """
    Build all export formats in one shot.
    """
    md = export_markdown(essay_text)
    txt = export_txt(essay_text)
    docx = export_docx(title, essay_text)
    pdf = export_pdf(title, essay_text)
    return ExportBundle(md=md, txt=txt, docx=docx, pdf=pdf)
